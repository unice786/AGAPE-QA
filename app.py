import os
import re
import pandas as pd
import customtkinter as ctk
from tkinter import filedialog, messagebox
from pypdf import PdfReader
import docx

# ------------------------------------------------------------
#  CONFIGURACIÓN INICIAL
# ------------------------------------------------------------
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

app = ctk.CTk()
app.geometry("750x700")
app.title("Project.D - Auditor Automático de Documentos")

ruta_carpeta = ctk.StringVar(value="Carpeta de tomos: No seleccionada")
ruta_listado = ctk.StringVar(value="Listado maestro F.GI.030: No seleccionado")

# ------------------------------------------------------------
#  FUNCIONES DE EXTRACCIÓN (MEJORADAS)
# ------------------------------------------------------------
def extraer_texto_pdf(ruta):
    """Extrae el texto de la primera página del PDF con manejo robusto de errores."""
    try:
        reader = PdfReader(ruta)
        if reader.is_encrypted:
            return "PDF PROTEGIDO CON CONTRASEÑA"
        if len(reader.pages) > 0:
            texto = reader.pages[0].extract_text()
            return texto if texto else "SIN TEXTO EXTRAÍBLE"
        return "PDF SIN PÁGINAS"
    except Exception as e:
        return f"ERROR AL LEER PDF: {str(e)[:100]}"

def extraer_texto_docx(ruta):
    """Extrae texto del cuerpo Y de los ENCABEZADOS/PIES DE PÁGINA reales de Word."""
    try:
        doc = docx.Document(ruta)
        texto = []
        # 1. Encabezados y pies de todas las secciones
        for seccion in doc.sections:
            # Encabezado
            if seccion.header:
                for p in seccion.header.paragraphs:
                    if p.text.strip():
                        texto.append(p.text)
            # Pie de página
            if seccion.footer:
                for p in seccion.footer.paragraphs:
                    if p.text.strip():
                        texto.append(p.text)
        # 2. Primeros párrafos del cuerpo (para tablas de control)
        for p in doc.paragraphs[:15]:
            texto.append(p.text)
        # 3. Primeras tablas (allí suelen estar los datos de revisión)
        for tabla in doc.tables[:3]:
            for fila in tabla.rows:
                for celda in fila.cells:
                    texto.append(celda.text)
        return "\n".join([t for t in texto if t])
    except Exception as e:
        return f"ERROR AL LEER DOCX: {str(e)[:100]}"

# ------------------------------------------------------------
#  REGEX CORREGIDAS (PARCHE EN ESPAÑOL Y MAYÚSCULAS)
# ------------------------------------------------------------
def buscar_metadatos(texto_interno):
    """
    Busca versión y fechas con patrones específicos del contexto institucional.
    Se limpian saltos de línea y se normaliza a mayúsculas.
    """
    if not texto_interno or "ERROR" in texto_interno.upper():
        return "NO DETECTADA", "NO DETECTADA", texto_interno[:250] if texto_interno else "SIN TEXTO"

    texto = re.sub(r'\s+', ' ', texto_interno).upper()

    # --- VERSIÓN ---
    version = "NO DETECTADA"
    patron_version = re.compile(
        r"""
        (?:
            REV(?:ISI[ÓO]N)?   # REV, REVISIÓN
            |VERSI[ÓO]N        # VERSIÓN
            |EDICI[ÓO]N        # EDICIÓN
            |ED\.?             # ED.
            |VER\.?\s*SI[ÓO]N  
        )
        [\s:\.\-]* # separadores
        ([0-9]+(维护)?(?:\.[0-9]+)?|[A-Z]{1,2})  # número o letras
        """,
        re.VERBOSE | re.IGNORECASE
    )
    
    match_v = patron_version.search(texto)
    if match_v:
        posible_version = match_v.group(1).strip()
        # PARCHE: Si capturó conectores en español como "DE", "DEL", "POR", se ignora
        if posible_version not in ["DE", "DEL", "POR", "EL", "LA"]:
            version = posible_version

    # --- FECHAS ---
    fechas_str = "NO DETECTADA"
    # PARCHE: Se cambió [a-záéíóú] por [A-ZÁÉÍÓÚ] para que machee con los meses en MAYÚSCULAS
    patron_fecha = re.findall(
        r'\b(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})\b|\b(\d{1,2}\s+DE\s+[A-ZÁÉÍÓÚ]+\s+DE\s+\d{2,4})\b',
        texto
    )
    
    fechas = []
    for tup in patron_fecha:
        if tup[0]:
            fechas.append(tup[0])
        elif tup[1]:
            fechas.append(tup[1])
            
    if fechas:
        fechas_str = " | ".join(fechas[:5])
    else:
        fechas_sueltas = re.findall(r'(\d{2}[\s]?\d{2}[\s]?\d{4})', texto)
        if fechas_sueltas:
            fechas_str = " | ".join(fechas_sueltas[:5])

    cabecera = texto[:250]

    return version, fechas_str, cabecera
# ------------------------------------------------------------
#  FUNCIONES DE INTERFAZ
# ------------------------------------------------------------
def seleccionar_carpeta():
    carpeta = filedialog.askdirectory(title="Selecciona la carpeta principal de los tomos")
    if carpeta:
        ruta_carpeta.set(carpeta)
        btn_iniciar.configure(state="normal", fg_color="#1f538d")
        imprimir_en_log(f"📁 Carpeta cargada: {carpeta}\nListo para extraer metadatos.\n", limpiar=True)

def seleccionar_listado():
    archivo = filedialog.askopenfilename(
        title="Selecciona el archivo Excel maestro (F.GI.030)",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    if archivo:
        ruta_listado.set(archivo)
        imprimir_en_log(f"📄 Listado maestro cargado: {os.path.basename(archivo)}\n")

def imprimir_en_log(texto, limpiar=False):
    caja_log.configure(state="normal")
    if limpiar:
        caja_log.delete("1.0", "end")
    caja_log.insert("end", texto)
    caja_log.see("end")
    caja_log.configure(state="disabled")
    app.update()

# ------------------------------------------------------------
#  CONCILIACIÓN CONTRA LISTADO MAESTRO (ESCÁNER DINÁMICO)
# ------------------------------------------------------------
def conciliar_con_listado(df_extraido, ruta_excel):
    """
    Carga el listado oficial, detecta automáticamente en qué fila están los 
    encabezados reales saltándose banners institucionales, y ejecuta el cruce.
    """
    try:
        # Leemos en bruto sin asumir filas de encabezado (header=None)
        df_crudo = pd.read_excel(ruta_excel, sheet_name=0, header=None)
    except Exception as e:
        imprimir_en_log(f"❌ No se pudo cargar el listado maestro: {e}\n")
        return df_extraido, None

    # 1. RASTREADOR: Buscar en qué fila vive la tabla real
    fila_encabezado = None
    for idx, fila in df_crudo.iterrows():
        valores_fila = [str(val).upper() for val in fila.values if pd.notna(val)]
        # Si la fila contiene palabras clave del F.GI.030, encontramos el inicio
        if any('CODIGO' in v or 'CÓDIGO' in v for v in valores_fila) and any('NOMBRE' in v or 'DOCUMENTO' in v for v in valores_fila):
            fila_encabezado = idx
            break

    if fila_encabezado is None:
        imprimir_en_log("❌ Error Estructural: No se encontró la fila con los encabezados 'CÓDIGO' y 'NOMBRE DEL DOCUMENTO'.\n")
        return df_extraido, None

    # 2. Reconstruir el DataFrame usando la fila correcta como nombres de columna
    columnas_reales = [str(c).strip() for c in df_crudo.iloc[fila_encabezado].tolist()]
    df_maestro = df_crudo.iloc[fila_encabezado + 1:].copy()
    df_maestro.columns = columnas_reales
    df_maestro = df_maestro.dropna(how='all') # Limpiar filas fantasma vacías al final

    # 3. Mapeo Flexible de Columnas detectadas
    col_codigo = None
    col_version_m = None
    col_fecha_m = None

    for col in df_maestro.columns:
        col_upper = str(col).upper()
        if 'CODIGO' in col_upper or 'CÓDIGO' in col_upper:
            col_codigo = col
        elif 'VERSION' in col_upper or 'VERSIÓN' in col_upper:
            col_version_m = col
        elif 'FECHA' in col_upper or 'VIGENCIA' in col_upper or 'ACTUALIZACIÓN' in col_upper:
            col_fecha_m = col

    if not col_codigo:
        imprimir_en_log("❌ No se pudo determinar con precisión la columna de Códigos en el maestro.\n")
        return df_extraido, None

    imprimir_en_log(f"🔍 Estructura F.GI.030 detectada con éxito en la fila {fila_encabezado + 1}.\n")
    imprimir_en_log(f"   • Buscando por la columna oficial: '{col_codigo}'\n")

    # 4. Preparar llaves de cruce limpias y en mayúsculas
    df_maestro = df_maestro.rename(columns={col_codigo: "Código Maestro"})
    df_maestro["Código Maestro"] = df_maestro["Código Maestro"].astype(str).str.upper().str.strip()

    df_extraido_limpio = df_extraido.copy()
    df_extraido_limpio["Código Exterior"] = df_extraido_limpio["Código Exterior"].astype(str).str.upper().str.strip()

    # 5. Cruce LEFT JOIN (Mantiene el control estricto: solo lo que existe físicamente en el tomo)
    df_merge = df_extraido_limpio.merge(
        df_maestro,
        left_on="Código Exterior",
        right_on="Código Maestro",
        how="left",
        indicator=True
    )

    df_merge["Estado en Listado"] = df_merge["_merge"].map({
        "both": "Registrado Oficialmente",
        "left_only": "FANTASMA (No mapeado en F.GI.030)"
    })

    # Normalizador de versiones para que "01" coincida con "1" o "REV 01"
    def normalizar_version(v):
        if pd.isna(v): return "SINDATO"
        v_limpia = re.sub(r'[^0-9A-Z]', '', str(v).upper().replace("REV", "").replace("VER", ""))
        return v_limpia.lstrip('0') if v_limpia.isdigit() else v_limpia

    # 6. Verificación cruzada inteligente de datos internos
    if col_version_m:
        df_merge["Versión en Listado"] = df_merge[col_version_m].fillna("N/A")
        df_merge["¿Versión coincide?"] = df_merge.apply(
            lambda row: "SÍ" if row["Estado en Listado"] == "Registrado Oficialmente" and normalizar_version(row["Versión Interior"]) == normalizar_version(row["Versión en Listado"])
            else ("NO" if row["Estado en Listado"] == "Registrado Oficialmente" else "N/A"),
            axis=1
        )
    else:
        df_merge["Versión en Listado"] = "No encontrada"
        df_merge["¿Versión coincide?"] = "N/A"

    if col_fecha_m:
        df_merge["Fecha en Listado"] = df_merge[col_fecha_m].fillna("N/A")
        df_merge["¿Fecha coincide?"] = df_merge.apply(
            lambda row: "SÍ" if row["Estado en Listado"] == "Registrado Oficialmente" and (str(row["Fechas Internas (Aprox)"]).strip() in str(row["Fecha en Listado"]).strip() or str(row["Fecha en Listado"]).strip() in str(row["Fechas Internas (Aprox)"]).strip())
            else ("NO" if row["Estado en Listado"] == "Registrado Oficialmente" else "N/A"),
            axis=1
        )
    else:
        df_merge["Fecha en Listado"] = "No encontrada"
        df_merge["¿Fecha coincide?"] = "N/A"

    # Eliminamos columnas técnicas de control de Pandas antes de entregar el archivo
    df_merge = df_merge.drop(columns=["_merge", "Código Maestro"])

    return df_extraido, df_merge
# ------------------------------------------------------------
#  MOTOR PRINCIPAL DE EXTRACCIÓN
# ------------------------------------------------------------
def iniciar_extraccion():
    if not ruta_carpeta.get():
        messagebox.showwarning("Aviso", "Primero selecciona la carpeta de los tomos.")
        return

    btn_iniciar.configure(state="disabled", text="Extrayendo y Procesando...")
    imprimir_en_log("Iniciando minería de datos físicos... Esto tomará unos segundos.\n", limpiar=True)

    datos_extraidos = []
    archivos_procesados = 0
    archivos_basura_ignorados = 0

    for ruta_actual, _, archivos in os.walk(ruta_carpeta.get()):
        for archivo in archivos:
            # FILTRO ANTI-BASURA
            if archivo.startswith('~$') or archivo.startswith('.') or archivo.endswith('.tmp'):
                archivos_basura_ignorados += 1
                continue

            archivos_procesados += 1
            tomo = os.path.basename(ruta_actual)
            nombre_base, ext = os.path.splitext(archivo)
            ext = ext.lower().strip()

            partes = nombre_base.strip().split(" ", 1)
            codigo_ext = partes[0].upper()
            titulo_ext = partes[1] if len(partes) > 1 else "SIN TÍTULO EXTERIOR"

            # Leer según formato
            ruta_completa = os.path.join(ruta_actual, archivo)
            texto_bruto = ""
            if ext == ".pdf":
                texto_bruto = extraer_texto_pdf(ruta_completa)
            elif ext == ".docx":
                texto_bruto = extraer_texto_docx(ruta_completa)

            version_int, fechas_int, cabecera = buscar_metadatos(texto_bruto)

            # Formatos legacy
            if ext in [".vsd", ".doc", ".xls", ".xlsx"]:
                version_int = "REQUIERE ABRIR MANUAL"
                fechas_int = "FORMATO NO ESCANEABLE"
                cabecera = "FORMATO LEGACY / BINARIO"

            datos_extraidos.append({
                "Tomo (Ubicación)": tomo,
                "Nombre Archivo Completo": archivo,
                "Código Exterior": codigo_ext,
                "Título Exterior": titulo_ext,
                "Versión Interior": version_int,
                "Fechas Internas (Aprox)": fechas_int,
                "Cabecera del Documento": cabecera,
                "Extensión": ext
            })

            if archivos_procesados % 100 == 0:
                imprimir_en_log(f"  ... {archivos_procesados} archivos reales analizados ...\n")

    imprimir_en_log(f"\n✔ Minería terminada.\n")
    imprimir_en_log(f"  • Archivos válidos procesados: {archivos_procesados}\n")
    imprimir_en_log(f"  • Archivos temporales ignorados: {archivos_basura_ignorados}\n")
    imprimir_en_log("-" * 60 + "\n")

    # Crear DataFrame base
    df_extraido = pd.DataFrame(datos_extraidos)

    # Si el usuario proporcionó el listado maestro, hacemos la conciliación
    df_conciliacion = None
    if ruta_listado.get() and os.path.exists(ruta_listado.get()):
        imprimir_en_log("Realizando cruce con el listado maestro F.GI.030...\n")
        df_extraido, df_conciliacion = conciliar_con_listado(df_extraido, ruta_listado.get())
        if df_conciliacion is not None:
            imprimir_en_log("✔ Cruce completado. Se generarán dos hojas en el Excel.\n")
        else:
            imprimir_en_log("⚠ No se pudo completar el cruce. Se exportará solo la extracción.\n")
    else:
        imprimir_en_log("ℹ No se seleccionó listado maestro. Se exportará solo la extracción.\n")

    # Exportar a Excel
    imprimir_en_log("Generando archivo de Excel de salida...\n")
    ruta_salida = os.path.join(ruta_carpeta.get(), "Reporte_Data_Extraida.xlsx")

    try:
        with pd.ExcelWriter(ruta_salida, engine='openpyxl') as writer:
            df_extraido.to_excel(writer, sheet_name="Datos_Extraidos", index=False)
            if df_conciliacion is not None:
                df_conciliacion.to_excel(writer, sheet_name="Conciliacion_FGI030", index=False)
        imprimir_en_log(f"🎉 ¡ÉXITO! El archivo está en:\n{ruta_salida}\n")
    except Exception as e:
        imprimir_en_log(f"❌ Error al guardar el Excel: {e}\n")

    btn_iniciar.configure(state="normal", text="Volver a Ejecutar Extracción")

# ------------------------------------------------------------
#  INTERFAZ GRÁFICA
# ------------------------------------------------------------
label_titulo = ctk.CTkLabel(app, text="Auditor Automático - Project.D", font=("Roboto", 22, "bold"))
label_titulo.pack(pady=20)

frame_seleccion = ctk.CTkFrame(app)
frame_seleccion.pack(pady=5, fill="x", padx=20)

btn_buscar = ctk.CTkButton(frame_seleccion, text="1. Seleccionar Carpeta USB", command=seleccionar_carpeta,
                           fg_color="#d35400", hover_color="#e67e22", width=200)
btn_buscar.grid(row=0, column=0, padx=10, pady=5)

label_ruta = ctk.CTkLabel(frame_seleccion, textvariable=ruta_carpeta, font=("Roboto", 10, "italic"), wraplength=400)
label_ruta.grid(row=0, column=1, padx=10, pady=5, sticky="w")

btn_listado = ctk.CTkButton(frame_seleccion, text="2. (Opcional) Seleccionar F.GI.030", command=seleccionar_listado,
                            fg_color="#2c3e50", hover_color="#34495e", width=200)
btn_listado.grid(row=1, column=0, padx=10, pady=5)

label_listado = ctk.CTkLabel(frame_seleccion, textvariable=ruta_listado, font=("Roboto", 10, "italic"), wraplength=400)
label_listado.grid(row=1, column=1, padx=10, pady=5, sticky="w")

caja_log = ctk.CTkTextbox(app, width=700, height=280, font=("Consolas", 11))
caja_log.pack(pady=10)
caja_log.insert("end", "Selecciona la carpeta de tomos y (opcional) el listado maestro para comenzar...\n")
caja_log.configure(state="disabled")

btn_iniciar = ctk.CTkButton(app, text="3. Iniciar Extracción y Generar Excel", command=iniciar_extraccion,
                            state="disabled", fg_color="#555555", height=45, font=("Roboto", 14, "bold"))
btn_iniciar.pack(pady=15)

if __name__ == "__main__":
    app.mainloop()