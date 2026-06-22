import os
import re
import subprocess
import tkinter as tk
from tkinter import filedialog, scrolledtext

try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

# ------------------------------------------------------------
#  CEREBRO DE EXTRACCIÓN OPTIMIZADO PARA DOCS / PDF (v2.4.0)
# ------------------------------------------------------------
def analizar_y_aislar_metadatos_docs(texto_completo, nombre_fichero=None):
    """
    Analizador de metadatos especializado para archivos de texto.
    Integra un sistema de validación cruzada con el nombre del fichero para PDFs.
    """
    # === 1. CONTROL DE CÓDIGO INTELIGENTE ===
    codigo_fichero = None
    if nombre_fichero:
        match_fn = re.search(r'([A-Z]\.[A-Z]{1,4}\.\d{3})', nombre_fichero.upper())
        if match_fn:
            codigo_fichero = match_fn.group(1)

    texto_comprimido = re.sub(r'[\s|]', '', texto_completo).upper()
    texto_comprimido = (texto_comprimido
                        .replace('Ó', 'O').replace('É', 'E')
                        .replace('Á', 'A').replace('Í', 'I').replace('Ú', 'U'))
    
    texto_comprimido_seguro = texto_comprimido.replace("FECHADEREVISION", "FECHA_REV").replace("FECHAREVISION", "FECHA_REV")

    if codigo_fichero:
        codigo_detectado = codigo_fichero
    else:
        match_cod = re.search(r'CODIGO[:\s|]*([A-Z]\.[A-Z]{1,4}\.\d{3})', texto_completo.upper())
        codigo_detectado = "NO DETECTADO"
        if match_cod:
            codigo_detectado = match_cod.group(1)
        else:
            match_suelto = re.search(r'([A-Z]\.[A-Z]{1,4}\.\d{3})', texto_comprimido_seguro)
            if match_suelto: 
                codigo_detectado = match_suelto.group(1)

    # === 2. EXTRACCIÓN DE TÍTULO ===
    patron_titulo = r'(?i)t[íi]tulo\s*[:\s|]\s*([\s\S]+?)(?=\s*(?:naturaleza|c[oó]digo|versi[oó]n|fecha|área|proceso|colegio)\s*[:|]|$)'
    match_tit = re.search(patron_titulo, texto_completo)
    titulo_detectado = "NO DETECTADO"
    
    if match_tit:
        raw_tit = match_tit.group(1).strip()
        raw_tit = re.split(r'(?i)\s*(?:[:|]\s*)?(?:naturaleza|c[oó]digo|versi[oó]n|fecha|área|proceso|colegio)', raw_tit)[0].strip()
        titulo_detectado = re.sub(r'\s*\|\s*', ' ', raw_tit).strip().strip('|').strip()

    # 🛡️ BLINDAJE ANTIFALLO PARA PDF: Si el título extraído es erróneo o se mezcló con etiquetas de control
    if nombre_fichero:
        nombre_sin_ext = os.path.splitext(nombre_fichero)[0]
        # Remueve el código del frente (ej: D.CS.001 ) para aislar el nombre puro del procedimiento
        titulo_fichero = re.sub(r'^[A-Z]\.[A-Z]{1,4}\.\d{3}\s*', '', nombre_sin_ext).strip()
        
        if (titulo_detectado == "NO DETECTADO" or 
            "NATURALEZA" in titulo_detectado.upper() or 
            "REVISION" in titulo_detectado.upper() or 
            len(titulo_detectado) <= 3):
            if titulo_fichero:
                titulo_detectado = titulo_fichero.upper()

    # === 3. NATURALEZA ===
    patron_naturaleza = r'(?i)naturaleza(?:[^|\n:]*)[:\s|]\s*([\s\S]+?)(?=\s*(?:t[íi]tulo|c[oó]digo|versi[oó]n|fecha)\s*[:|]|$)'
    match_nat = re.search(patron_naturaleza, texto_completo)
    naturaleza_detectada = "NO DETECTADA"
    if match_nat:
        raw_nat = match_nat.group(1).strip()
        raw_nat = re.split(r'(?i)\s*(?:[:|]\s*)?(?:t[íi]tulo|c[oó]digo|versi[oó]n|fecha)', raw_nat)[0].strip()
        raw_nat = re.sub(r'\s*\|\s*', ' ', raw_nat).strip().strip('|').strip()
        if raw_nat and not raw_nat.upper().startswith("TITULO"):
            naturaleza_detectada = raw_nat

    # === 4. EDICIÓN / VERSIÓN ===
    patron_version = r'(?i)(?:versi[oó]n|revisi[oó]n|rev)\s*[:\s|]\s*(\d+)'
    match_ver = re.search(patron_version, texto_completo)
    version_detectada = "NO NOTADA"
    if match_ver:
        version_detectada = match_ver.group(1)
        if len(version_detectada) == 1: 
            version_detectada = f"0{version_detectada}"
    else:
        texto_sin_fechas = re.sub(r'\d{1,2}/\d{1,2}/\d{2,4}', '', texto_comprimido_seguro)
        match_ver_comp = re.search(r'(?:VERSION|REVISION|REV):?(\d+)', texto_sin_fechas)
        if match_ver_comp:
            version_detectada = match_ver_comp.group(1)
            if len(version_detectada) == 1: version_detectada = f"0{version_detectada}"

    # === 5. EXTRACCIÓN DE FECHAS ===
    fechas_encontradas = re.findall(r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', texto_completo)
    fecha_ini_detectada = "NO DETECTADA"
    fecha_rev_detectada = "NO DETECTADA"
   
    if len(fechas_encontradas) >= 1:
        def clave_cronologica(f_str):
            try:
                partes = f_str.split('/')
                dia = int(partes[0])
                mes = int(partes[1])
                anio = int(partes[2])
                if anio < 100: anio += 2000
                return (anio, mes, dia)
            except:
                return (0, 0, 0)
        
        fechas_ordenadas = sorted(list(set(fechas_encontradas)), key=clave_cronologica)
        if len(fechas_ordenadas) >= 1:
            fecha_ini_detectada = fechas_ordenadas[0]
        if len(fechas_ordenadas) >= 2:
            fecha_rev_detectada = fechas_ordenadas[-1]

    return titulo_detectado, codigo_detectado, version_detectada, fecha_ini_detectada, fecha_rev_detectada, naturaleza_detectada

# ------------------------------------------------------------
#  MOTORES DE LECTURA DE ARCHIVOS DE TEXTO
# ------------------------------------------------------------
def extraer_todo_documento(ruta):
    ext = os.path.splitext(ruta)[1].lower()
    acumulado = []
    
    if ext == '.docx' and docx is not None:
        try:
            doc = docx.Document(ruta)
            for p in doc.paragraphs:
                if p.text.strip(): acumulado.append(p.text.strip())
            for t in doc.tables:
                for f in t.rows:
                    for c in f.cells:
                        if c.text.strip(): acumulado.append(c.text.strip())
            return " \n ".join(acumulado)
        except Exception as e:
            return f"ERROR leyendo DOCX: {str(e)}"
            
    elif ext == '.doc':
        try:
            carpeta_destino = "/tmp"
            comando = ['libreoffice', '--headless', '--convert-to', 'docx', ruta, '--outdir', carpeta_destino]
            subprocess.run(comando, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
            
            nombre_base = os.path.splitext(os.path.basename(ruta))[0]
            ruta_docx_temp = os.path.join(carpeta_destino, nombre_base + '.docx')
            
            if os.path.exists(ruta_docx_temp):
                doc = docx.Document(ruta_docx_temp)
                for p in doc.paragraphs:
                    if p.text.strip(): acumulado.append(p.text.strip())
                for t in doc.tables:
                    for f in t.rows:
                        for c in f.cells:
                            if c.text.strip(): acumulado.append(c.text.strip())
                os.remove(ruta_docx_temp)
                return " \n ".join(acumulado)
        except Exception as e:
            return f"ERROR en conversión interna de .doc: {str(e)}"

    elif ext == '.pdf' and pypdf is not None:
        try:
            reader = pypdf.PdfReader(ruta)
            for pagina in reader.pages:
                txt = pagina.extract_text()
                if txt: acumulado.append(txt)
            return " \n ".join(acumulado)
        except Exception as e:
            return f"ERROR leyendo PDF: {str(e)}"
            
    return "(Formato no soportado)"

# ------------------------------------------------------------
#  CONTROLADOR DE LA INTERFAZ GUI
# ------------------------------------------------------------
def abrir_y_analizar():
    ruta_archivo = filedialog.askopenfilename(
        title="Cargar Documento para Auditoría",
        filetypes=[("Documentos de Texto", "*.docx *.doc *.pdf")]
    )
    if not ruta_archivo: return

    nombre_fichero = os.path.basename(ruta_archivo)
    texto_crudo = extraer_todo_documento(ruta_archivo)

    # Ejecución del cerebro con el validador cruzado de nombre activado
    titulo, codigo, version, fecha_ini, fecha_rev, naturaleza = analizar_y_aislar_metadatos_docs(texto_crudo, nombre_fichero)
   
    lbl_archivo.config(text=f"Fichero: {nombre_fichero}", fg="#00f5d4")
    txt_visor.delete("1.0", tk.END)
   
    txt_visor.insert(tk.END, f"📄 AGAPE QA - EXTRACTOR MULTI-FORMATO AVANZADO (v2.4.0-pdf_fix)\n")
    txt_visor.insert(tk.END, f"===========================================================\n")
    txt_visor.insert(tk.END, f"📁 ARCHIVO EVALUADO: {nombre_fichero}\n\n")
   
    txt_visor.insert(tk.END, f"📌 [DATOS AISLADOS DEL ENCABEZADO REAL]:\n")
    txt_visor.insert(tk.END, f"   ➔ TÍTULO INTERNO:      {titulo}\n")
    txt_visor.insert(tk.END, f"   ➔ NATURALEZA DOC:      {naturaleza}\n")
    txt_visor.insert(tk.END, f"   ➔ CÓDIGO DEL FORMATO:  {codigo}\n")
    txt_visor.insert(tk.END, f"   ➔ EDICIÓN / REVISIÓN:  {version}\n")
    txt_visor.insert(tk.END, f"   ➔ FECHA INICIAL:       {fecha_ini}\n")
    txt_visor.insert(tk.END, f"   ➔ FECHA DE REVISIÓN:   {fecha_rev}\n")
    txt_visor.insert(tk.END, f"===========================================================\n\n")
   
    txt_visor.insert(tk.END, f"[📋 FLUJO TEXTUAL EVALUADO]:\n")
    txt_visor.insert(tk.END, texto_crudo if texto_crudo.strip() else "(Documento sin texto legible)")

# ------------------------------------------------------------
#  CONFIGURACIÓN DE LA INTERFAZ DE USUARIO
# ------------------------------------------------------------
ventana = tk.Tk()
ventana.title("AGAPE QA - Extractor de Documentos (.docx / .doc / .pdf)")
ventana.geometry("950x720") 

lbl_titulo = tk.Label(ventana, text="Fase 1: Extractor Avanzado de Metadatos (.docx / .doc / .pdf)", font=("Arial", 12, "bold"))
lbl_titulo.pack(pady=10)

btn_subir = tk.Button(
    ventana,
    text="📁 Cargar Documento (.docx / .doc / .pdf)",
    font=("Arial", 11, "bold"),
    command=abrir_y_analizar,
    bg="#24a0ed",
    fg="white",
    padx=20,
    pady=6
)
btn_subir.pack(pady=5)

lbl_archivo = tk.Label(ventana, text="Ningún documento seleccionado", font=("Arial", 10, "italic"), fg="gray")
lbl_archivo.pack(pady=5)

# Pantalla negra terminal con la clásica letra verde fosforescente (#39FF14)
txt_visor = scrolledtext.ScrolledText(ventana, wrap=tk.WORD, font=("Courier New", 10), bg="#0f0f1c", fg="#39FF14")
txt_visor.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

ventana.mainloop()
